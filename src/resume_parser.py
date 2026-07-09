"""
src/resume_parser.py

CV/resume intake: extract text from PDF/DOCX/TXT, structure it via the
shared call_llm() pipeline (Ollama -> Groq fallback), and validate
extracted skills against skills_taxonomy.csv.

Never raises on bad input -- always returns a dict with an "error" key
set on failure, so the dashboard (Phase 4) can render a friendly
message instead of crashing on a malformed upload.

availability_pct and proficiency levels are NOT extracted -- they are
business inputs a resume can't tell you, and are left blank for the
human review step (see plan item 4 in the phase list). Never auto-fill
these with guesses.
"""

import json
import re
from pathlib import Path

import pandas as pd

from generate_explanation import call_llm

# ── Text extraction ─────────────────────────────────────────────────

def extract_text_from_file(file_path: str) -> str:
    """
    Extract raw text from a resume file. Supports .pdf, .docx, .txt.
    Raises ValueError for unsupported extensions or empty extraction --
    caller (parse_resume) catches this and turns it into an error dict.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        text = _extract_pdf(path)
    elif ext == ".docx":
        text = _extract_docx(path)
    elif ext == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: .pdf, .docx, .txt"
        )

    text = text.strip()
    if not text:
        raise ValueError(
            "No extractable text found -- file may be a scanned image "
            "PDF with no text layer, or empty."
        )
    return text


def _extract_pdf(path: Path) -> str:
    import pdfplumber
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                chunks.append(page_text)
    return "\n".join(chunks)


def _extract_docx(path: Path) -> str:
    import docx
    doc = docx.Document(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull table cells -- resumes sometimes put skills in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paras.append(cell.text.strip())
    return "\n".join(paras)


# ── LLM structuring ─────────────────────────────────────────────────

EXTRACTION_PROMPT_TEMPLATE = """You are extracting structured data from a resume for an internal staffing system. Read the resume text below and return ONLY a JSON object -- no preamble, no markdown code fences, no explanation.

Return exactly this shape:
{{
  "name": "string, the candidate's full name",
  "role": "string, their primary job title / role (e.g. 'Backend Dev', 'Data Scientist')",
  "experience_years": integer, total years of professional experience (best estimate, use 0 if unclear),
  "skills": ["list", "of", "individual", "skill", "strings"],
  "department": "string, best-guess department (e.g. 'Engineering', 'Data', 'Product') or null if unclear",
  "location": "string, city or 'Remote' if stated, or null if not mentioned"
}}

If a field genuinely cannot be determined from the text, use null (or 0 for experience_years, or [] for skills) -- do not invent values.

Resume text:
---
{resume_text}
---

JSON:"""


def build_extraction_prompt(resume_text: str, max_chars: int = 6000) -> str:
    """Truncate very long resumes to keep the prompt within a sane
    token budget -- most resumes are 1-3 pages and won't hit this."""
    if len(resume_text) > max_chars:
        resume_text = resume_text[:max_chars] + "\n...[truncated]"
    return EXTRACTION_PROMPT_TEMPLATE.format(resume_text=resume_text)


def _parse_llm_json(raw: str) -> dict:
    """
    LLMs sometimes wrap JSON in ```json fences or add stray text
    despite instructions. Strip fences, then find the outermost
    {...} block and parse that.
    """
    cleaned = re.sub(r"```(?:json)?", "", raw).strip()
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end < start:
        raise ValueError("No JSON object found in LLM response")
    return json.loads(cleaned[start:end + 1])


# ── Skill taxonomy validation ───────────────────────────────────────

def _load_taxonomy_skills(taxonomy_path: str) -> set:
    """
    Loads skills_taxonomy.csv and returns a lowercase set of known
    skill names. Auto-detects the column name since I haven't verified
    your exact schema -- checks common names first, falls back to the
    first column. Returns an empty set (never raises) if the file is
    missing, so validation degrades to "nothing flagged" rather than
    crashing the whole parse.
    """
    path = Path(taxonomy_path)
    if not path.exists():
        return set()

    df = pd.read_csv(path)
    for candidate_col in ("skill", "skill_name", "name"):
        if candidate_col in df.columns:
            col = candidate_col
            break
    else:
        col = df.columns[0]

    return set(df[col].astype(str).str.strip().str.lower())


def validate_skills(extracted_skills: list,
                     taxonomy_path: str = "data/processed/skills_taxonomy.csv") -> dict:
    """
    Splits extracted skills into recognized vs unrecognized against
    the taxonomy. Soft validation only -- unrecognized skills are
    flagged for the review UI, never silently dropped or auto-corrected.
    """
    known = _load_taxonomy_skills(taxonomy_path)
    matched, unrecognized = [], []

    for skill in extracted_skills:
        if not skill or not str(skill).strip():
            continue
        (matched if str(skill).strip().lower() in known else unrecognized).append(
            str(skill).strip()
        )

    return {
        "matched": matched,
        "unrecognized": unrecognized,
        "taxonomy_loaded": len(known) > 0,
    }

# ── Duplicate detection ──────────────────────────────────────────────

import difflib


def _normalize_name(name) -> str:
    if not name:
        return ""
    return re.sub(r"\s+", " ", str(name).strip().lower())


def find_possible_duplicates(candidate: dict,
                              employees_df: pd.DataFrame,
                              custom_employees: list = None,
                              name_col: str = "name",
                              role_col: str = "role",
                              dept_col: str = "department",
                              fuzzy_threshold: float = 0.88) -> list:
    """
    Soft duplicate check against both the real employee roster and
    session-added custom employees. Never blocks -- returns matches
    for the caller (dashboard review UI) to show as a warning with an
    explicit "add anyway" override, per the plan: real names collide
    legitimately, so this must not hard-block.

    Two match tiers:
      - "exact": normalized name is identical (same as the original
        design sketch -- case/whitespace-insensitive equality).
      - "fuzzy": name similarity >= fuzzy_threshold but not exact --
        catches likely typo re-entries of the same person (e.g.
        "Adiiti Sharma" vs "Aditi Sharma"). Kept as a separate, lower-
        confidence tier so the review UI can label it differently
        ("possible variation" vs "exact match") rather than treating
        both the same.

    `candidate` is expected to have at least a "name" key (e.g. the
    dict returned by parse_resume(), or a manually-filled form dict --
    both paths use the same review form per the plan, so this function
    doesn't care which path produced the candidate).

    Returns a list of dicts, sorted by similarity descending:
      {employee_id, name, role, department, source, match_type, similarity}
    "source" is "existing" or "custom". If employee_id isn't present
    on a custom-employee dict yet (not assigned until commit), it's
    omitted from that match's dict rather than faked.
    """
    candidate_norm = _normalize_name(candidate.get("name"))
    if not candidate_norm:
        return []  # can't compare without a name -- nothing to flag

    matches = []

    def _consider(emp_name, emp_role, emp_dept, emp_id, source):
        emp_norm = _normalize_name(emp_name)
        if not emp_norm:
            return
        if emp_norm == candidate_norm:
            match_type, similarity = "exact", 1.0
        else:
            similarity = difflib.SequenceMatcher(None, candidate_norm, emp_norm).ratio()
            if similarity < fuzzy_threshold:
                return
            match_type = "fuzzy"

        entry = {
            "name": emp_name,
            "role": emp_role,
            "department": emp_dept,
            "source": source,
            "match_type": match_type,
            "similarity": round(similarity, 3),
        }
        if emp_id is not None:
            entry["employee_id"] = emp_id
        matches.append(entry)

    # Real employee roster
    if employees_df is not None and not employees_df.empty and name_col in employees_df.columns:
        id_col = employees_df.index.name  # matcher.py convention: set_index("employee_id")
        for idx, row in employees_df.iterrows():
            _consider(
                emp_name=row.get(name_col),
                emp_role=row.get(role_col) if role_col in employees_df.columns else None,
                emp_dept=row.get(dept_col) if dept_col in employees_df.columns else None,
                emp_id=idx if id_col == "employee_id" else row.get("employee_id"),
                source="existing",
            )

    # Session-added custom employees (list of dicts, e.g. st.session_state.custom_employees)
    for emp in (custom_employees or []):
        _consider(
            emp_name=emp.get(name_col, emp.get("name")),
            emp_role=emp.get(role_col, emp.get("role")),
            emp_dept=emp.get(dept_col, emp.get("department")),
            emp_id=emp.get("employee_id"),
            source="custom",
        )

    matches.sort(key=lambda m: m["similarity"], reverse=True)
    return matches


def format_duplicate_warning(matches: list) -> list:
    """
    Turns find_possible_duplicates() output into display-ready
    strings for the dashboard, matching the format from the design
    doc: "E014 — Aryan Maharaj, Frontend Dev, Bangalore already
    exists — is this the same person?"
    Fuzzy matches get a softer phrasing since they're a weaker signal.
    """
    lines = []
    for m in matches:
        label = m.get("employee_id", "(unsaved)")
        role = m.get("role") or "role unknown"
        dept = m.get("department") or "dept unknown"
        if m["match_type"] == "exact":
            lines.append(
                f"{label} — {m['name']}, {role}, {dept} already exists "
                f"— is this the same person?"
            )
        else:
            lines.append(
                f"{label} — {m['name']}, {role}, {dept} looks similar "
                f"({int(m['similarity']*100)}% name match) — possible "
                f"duplicate or typo?"
            )
    return lines

# ── Main entry point ────────────────────────────────────────────────

def parse_resume(file_path: str,
                  taxonomy_path: str = "data/processed/skills_taxonomy.csv") -> dict:
    """
    Full pipeline: extract text -> LLM structuring -> skill validation.

    Returns a dict always containing a "backend" key ("ollama"/"groq"/
    "error") so the dashboard can show provenance. On any failure,
    returns {"error": "...", "backend": "error"} instead of raising --
    this MUST stay true through Phase 4, since a bad upload should
    never crash the Streamlit app.

    availability_pct, proficiency levels, and employee_id are
    deliberately absent from this return value -- those are assigned
    later, during the human review step, not by this function.
    """
    try:
        text = extract_text_from_file(file_path)
    except Exception as e:
        return {"error": f"Could not read file: {e}", "backend": "error"}

    prompt = build_extraction_prompt(text)
    raw_response, backend = call_llm(prompt, max_tokens=500)

    if backend == "error":
        return {"error": raw_response, "backend": "error"}

    try:
        extracted = _parse_llm_json(raw_response)
    except Exception as e:
        return {
            "error": f"LLM returned unparseable response: {e}",
            "backend": backend,
            "raw_response": raw_response,  # for debugging in the review UI
        }

    skills = extracted.get("skills") or []
    if not isinstance(skills, list):
        skills = [skills]

    validation = validate_skills(skills, taxonomy_path=taxonomy_path)

    return {
        "name": extracted.get("name"),
        "role": extracted.get("role"),
        "experience_years": extracted.get("experience_years", 0),
        "skills": skills,
        "department": extracted.get("department"),
        "location": extracted.get("location"),
        "skill_validation": validation,
        "backend": backend,
        "error": None,
    }


# ── Main ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    print("\n── Resume Parser Test ──────────────────────────────\n")

    if len(sys.argv) > 1:
        test_file = sys.argv[1]
    else:
        print("Usage: python resume_parser.py <path_to_resume.pdf|.docx|.txt>")
        sys.exit(1)

    result = parse_resume(test_file)

    if result.get("error"):
        print(f"   Error: {result['error']}")
    else:
        print(f"   Backend: {result['backend']}")
        print(f"   Name: {result['name']}")
        print(f"   Role: {result['role']}")
        print(f"   Experience: {result['experience_years']} years")
        print(f"   Department: {result['department']}")
        print(f"   Location: {result['location']}")
        print(f"   Skills extracted: {result['skills']}")
        v = result["skill_validation"]
        print(f"   Matched against taxonomy: {v['matched']}")
        print(f"   Unrecognized (flag for review): {v['unrecognized']}")
        if not v["taxonomy_loaded"]:
            print("   Note: skills_taxonomy.csv not found or empty -- "
                  "nothing could be validated, everything shown as "
                  "unrecognized by default.")

    print("\n── Done ─────────────────────────────────────────────")